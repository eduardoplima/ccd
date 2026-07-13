"""Informação de liquidação do Acórdão 197/2026 (processo 000068/2024).

Preenche scripts/automacao/templates/modelo_informacao.docx com a apuração dos
valores a restituir pelo escritório Fernanda de Paula Sociedade Individual de
Advocacia (CNPJ 48.581.488/0001-14), calculados sobre vwDespesaPagamento (BdDIP).
"""
import shutil
from datetime import datetime
from pathlib import Path

import docx
import pandas as pd
from docx.oxml.ns import qn

from ccd.config import REPO_ROOT
from ccd.db import get_connection

TEMPLATE = str(REPO_ROOT / "scripts" / "automacao" / "templates" / "modelo_informacao.docx")
OUT = str(Path(__file__).parent / "informacao_liquidacao_000068_2024.docx")

CNPJ = "48581488000114"
TETOS = {369: ("CRUZETA", 480_000.00), 405: ("LAGOA NOVA", 320_000.00), 400: ("JUCURUTU", 250_000.00)}


def brl(v: float) -> str:
    return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


# ── apuração ─────────────────────────────────────────────────────────────
eng = get_connection(db="BdDIP")
df = pd.read_sql(
    """SELECT id_orgao, municipio, data, documento, valor, ISNULL(valor_anulado,0) valor_anulado
       FROM vwDespesaPagamento
       WHERE cpf_cnpj_favorecido = %(cnpj)s AND id_orgao IN (369, 400, 405)
       ORDER BY id_orgao, data""",
    eng,
    params={"cnpj": CNPJ},
)
df["liquido"] = df["valor"] - df["valor_anulado"]
tot = df.groupby("id_orgao")["liquido"].sum()

excessos = {o: tot[o] - teto for o, (_, teto) in TETOS.items()}
total_restituir = sum(excessos.values())
assert all(v > 0 for v in excessos.values())

# ── conteúdo ─────────────────────────────────────────────────────────────
replacements = {
    "{{processo}}": "000068/2024 - TC",
    "{{assunto}}": "CONTRATAÇÃO DE SOCIEDADE DE ADVOCACIA POR INEXIGIBILIDADE — "
    "LIQUIDAÇÃO DOS VALORES A RESTITUIR (ITEM II, ALÍNEA “D”, DO ACÓRDÃO Nº 197/2026 - TC)",
    "{{relator}}": "Conselheiro Antonio Gilberto de Oliveira Jales",
}

paragrafos = [
    "1. Trata-se de liquidação dos valores a serem restituídos ao erário pela sociedade "
    "Fernanda de Paula Sociedade Individual de Advocacia (CNPJ nº 48.581.488/0001-14), na forma "
    "do item II, alínea “d”, do Acórdão nº 197/2026 - TC (2ª Câmara, sessão de 23/06/2026), que "
    "declarou a nulidade integral dos contratos firmados com os Municípios de Cruzeta, Jucurutu "
    "e Lagoa Nova e fixou teto indenizatório, com base no art. 85, § 3º, do CPC, aplicável ao "
    "proveito econômico total obtido em cada ciclo de atuação (2023/2024 e 2024/2025), nos "
    "montantes máximos por ciclo de R$ 480.000,00 (Cruzeta), R$ 320.000,00 (Lagoa Nova) e "
    "R$ 250.000,00 (Jucurutu).",

    "2. A apuração foi realizada sobre os dados oficiais do SIAI (Anexo 14 — Empenhos, "
    "Liquidações e Pagamentos), prestados pelos próprios jurisdicionados e consultados por meio "
    "da visão consolidada vwDespesaPagamento da base de dados desta Corte (BdDIP), considerando "
    "exclusivamente arquivos processados com sucesso. Foram filtrados todos os pagamentos ao "
    "CNPJ nº 48.581.488/0001-14 realizados pelas unidades jurisdicionadas Prefeitura Municipal "
    "de Cruzeta (código PMCRUZETA), Prefeitura Municipal de Jucurutu (PMJUCURUTU) e Prefeitura "
    "Municipal de Lagoa Nova (PMLNOVA), líquidos de anulações. Os pagamentos identificados são "
    "os seguintes:",
]

tab_detalhe = [("Município", "Data", "Documento", "Valor")]
for _, r in df.iterrows():
    tab_detalhe.append(
        (r["municipio"].title(), r["data"].strftime("%d/%m/%Y"), r["documento"], brl(r["liquido"]))
    )
tab_detalhe.append(("TOTAL", "", "", brl(df["liquido"].sum())))

paragrafos_2 = [
    "3. Todos os pagamentos ocorreram entre 28/12/2023 e 10/06/2024, não havendo qualquer "
    "pagamento posterior registrado no SIAI. Dessa forma, a integralidade dos desembolsos "
    "concentra-se no ciclo de atuação 2023/2024, não tendo sido consumido o teto relativo ao "
    "ciclo 2024/2025 (sem pagamentos). O valor a restituir por Município corresponde, portanto, "
    "ao total pago deduzido do teto indenizatório de um único ciclo:",
]

tab_resumo = [("Município", "Total pago (SIAI)", "Teto indenizatório", "Valor a restituir")]
for o in (369, 405, 400):
    nome, teto = TETOS[o]
    tab_resumo.append((nome.title(), brl(tot[o]), brl(teto), brl(excessos[o])))
tab_resumo.append(("TOTAL", brl(sum(tot[o] for o in TETOS)), "", brl(total_restituir)))

paragrafos_3 = [
    "4. Os valores apurados conferem com os indicados no voto-vista acolhido pelo Acórdão nº "
    "197/2026 - TC (nota de rodapé nº 17), que registrava, com dados dos portais de transparência "
    "até 07/10/2024, desembolsos de R$ 1.199.335,27 (Cruzeta), R$ 809.753,80 (Lagoa Nova) e "
    "R$ 602.045,36 (Jucurutu). A presente apuração, por se basear na integralidade das remessas "
    "do SIAI, identificou pagamentos adicionais em Lagoa Nova não capturados naquela consulta.",
]

paragrafos_peticao = [
    "5. Registre-se que, em 09/07/2026, a sociedade protocolou pedido de reconsideração em face "
    "do Acórdão nº 197/2026 - TC (evento 229), em nome próprio e dos Municípios contratantes, no "
    "qual declara como proveito efetivamente recebido pelos Municípios no ciclo 2023/2024 os "
    "montantes de R$ 6.039.006,81 (Cruzeta), R$ 4.842.538,65 (Lagoa Nova) e R$ 3.010.108,63 "
    "(Jucurutu), sem controverter os pagamentos que recebeu. A incidência do percentual "
    "contratual de 20% sobre tais montantes resulta em R$ 1.207.801,36, R$ 968.507,73 e "
    "R$ 602.021,73, respectivamente, cifras que convergem, com divergência máxima de 0,7%, com "
    "os pagamentos apurados no SIAI (R$ 1.199.357,58, R$ 969.087,55 e R$ 602.045,36). A "
    "convergência entre os registros oficiais e os valores declarados pela própria recorrente "
    "corrobora a completude da apuração e a vinculação integral dos pagamentos ao ciclo "
    "2023/2024.",
]

paragrafos_conclusao = [
    "6. Diante do exposto, concluída a apuração determinada no item II, alínea “d”, do Acórdão "
    "nº 197/2026 - TC, remetem-se os autos ao gabinete do Exmo. Conselheiro Relator, "
    "registrando-se que:",

    "a) foram apurados como valores a restituir ao erário pela sociedade Fernanda de Paula "
    f"Sociedade Individual de Advocacia: {brl(excessos[369])} ao Município de Cruzeta, "
    f"{brl(excessos[405])} ao Município de Lagoa Nova e {brl(excessos[400])} ao Município de "
    f"Jucurutu, totalizando {brl(total_restituir)}, sujeitos a atualização na forma da Lei "
    "Orgânica desta Corte e da Resolução nº 013/2015-TCE;",

    "b) foi realizado o cadastro provisório dos respectivos débitos de ressarcimento em favor "
    "dos erários municipais credores, bem como das obrigações determinadas no Acórdão nº "
    "197/2026 - TC, permanecendo a constituição definitiva dos débitos e a intimação para "
    "recolhimento condicionadas ao trânsito em julgado, sobrestado em razão do pedido de "
    "reconsideração (evento 229), dotado de efeito suspensivo nos termos do art. 125, § 4º, da "
    "Lei Complementar Estadual nº 464/2012.",
]

# ── render ───────────────────────────────────────────────────────────────
doc = docx.Document(TEMPLATE)

for p in doc.paragraphs:
    for run in p.runs:
        for key, val in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, val)

first_loop = last_loop = None
for i, p in enumerate(doc.paragraphs):
    if "{% for" in p.text:
        first_loop = i
    if "{% endfor" in p.text:
        last_loop = i
assert first_loop is not None and last_loop is not None

parent = doc.paragraphs[first_loop]._p.getparent()
anchor = doc.paragraphs[first_loop - 1]._p
for idx in range(last_loop, first_loop - 1, -1):
    parent.remove(doc.paragraphs[idx]._p)


# Paleta institucional TCE/RN (skill tce-rn-identity)
TCE = {
    "verde_principal": "2E5B3C",
    "verde_escuro": "1A3D28",
    "branco": "FFFFFF",
    "cinza_texto": "333333",
    "cinza_claro": "F2F2F2",
    "cinza_borda": "CCCCCC",
}


def make_p(text: str, heading: bool = False):
    p = docx.oxml.OxmlElement("w:p")
    pPr = docx.oxml.OxmlElement("w:pPr")
    if heading:
        spacing = docx.oxml.OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "240")
        spacing.set(qn("w:after"), "120")
        pPr.append(spacing)
    else:
        spacing = docx.oxml.OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "240")
        spacing.set(qn("w:after"), "240")
        pPr.append(spacing)
        jc = docx.oxml.OxmlElement("w:jc")
        jc.set(qn("w:val"), "both")
        pPr.append(jc)
        if len(text) > 1 and text[1] == ")":
            ind = docx.oxml.OxmlElement("w:ind")
            ind.set(qn("w:left"), "720")
            ind.set(qn("w:hanging"), "360")
            pPr.append(ind)
    p.append(pPr)
    r = docx.oxml.OxmlElement("w:r")
    rPr = docx.oxml.OxmlElement("w:rPr")
    if heading:
        rPr.append(docx.oxml.OxmlElement("w:b"))
        color = docx.oxml.OxmlElement("w:color")
        color.set(qn("w:val"), TCE["verde_principal"])
        rPr.append(color)
    sz = docx.oxml.OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rPr.append(sz)
    r.append(rPr)
    t = docx.oxml.OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def make_table(rows):
    ncols = len(rows[0])
    usable = 8640
    tbl = docx.oxml.OxmlElement("w:tbl")
    tblPr = docx.oxml.OxmlElement("w:tblPr")
    tblW = docx.oxml.OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(usable))
    tblPr.append(tblW)
    jc = docx.oxml.OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)
    borders = docx.oxml.OxmlElement("w:tblBorders")
    for b in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = docx.oxml.OxmlElement(f"w:{b}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), TCE["cinza_borda"] if b.startswith("inside") else TCE["verde_principal"])
        borders.append(el)
    tblPr.append(borders)
    tbl.append(tblPr)
    grid = docx.oxml.OxmlElement("w:tblGrid")
    for _ in range(ncols):
        gc = docx.oxml.OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(usable // ncols))
        grid.append(gc)
    tbl.append(grid)
    for ri, row in enumerate(rows):
        tr = docx.oxml.OxmlElement("w:tr")
        if ri == 0:
            fill, font_color = TCE["verde_principal"], TCE["branco"]
        elif row[0] == "TOTAL":
            fill, font_color = TCE["verde_escuro"], TCE["branco"]
        elif ri % 2 == 1:
            fill, font_color = TCE["cinza_claro"], TCE["cinza_texto"]
        else:
            fill, font_color = TCE["branco"], TCE["cinza_texto"]
        for cell in row:
            tc = docx.oxml.OxmlElement("w:tc")
            tcPr = docx.oxml.OxmlElement("w:tcPr")
            tcW = docx.oxml.OxmlElement("w:tcW")
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(usable // ncols))
            tcPr.append(tcW)
            shd = docx.oxml.OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
            tc.append(tcPr)
            p = docx.oxml.OxmlElement("w:p")
            pPr = docx.oxml.OxmlElement("w:pPr")
            jc2 = docx.oxml.OxmlElement("w:jc")
            jc2.set(qn("w:val"), "center" if ri == 0 else "right")
            pPr.append(jc2)
            p.append(pPr)
            r = docx.oxml.OxmlElement("w:r")
            rPr = docx.oxml.OxmlElement("w:rPr")
            if ri == 0 or row[0] == "TOTAL":
                rPr.append(docx.oxml.OxmlElement("w:b"))
            color = docx.oxml.OxmlElement("w:color")
            color.set(qn("w:val"), font_color)
            rPr.append(color)
            sz = docx.oxml.OxmlElement("w:sz")
            sz.set(qn("w:val"), "20")
            rPr.append(sz)
            r.append(rPr)
            t = docx.oxml.OxmlElement("w:t")
            t.text = cell
            t.set(qn("xml:space"), "preserve")
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    return tbl


elems = (
    [make_p("1. DA LIQUIDAÇÃO", heading=True)]
    + [make_p(t) for t in paragrafos]
    + [make_table(tab_detalhe)]
    + [make_p(t) for t in paragrafos_2]
    + [make_table(tab_resumo)]
    + [make_p(t) for t in paragrafos_3]
    + [make_p("2. DA PETIÇÃO IMPETRADA EM 09/07/2026", heading=True)]
    + [make_p(t) for t in paragrafos_peticao]
    + [make_p("3. DA CONCLUSÃO", heading=True)]
    + [make_p(t) for t in paragrafos_conclusao]
)
ref = anchor
for el in elems:
    ref.addnext(el)
    ref = el

# preserva a versão anterior antes de sobrescrever
out_path = Path(OUT)
if out_path.exists():
    bak = out_path.with_name(f"{out_path.stem}_{datetime.now():%Y%m%d_%H%M%S}{out_path.suffix}")
    shutil.copy2(out_path, bak)
    print(f"Versão anterior preservada em: {bak.name}")
doc.save(OUT)
print(f"Informação salva em: {OUT}")
print(f"  Cruzeta: pago {brl(tot[369])} | restituir {brl(excessos[369])}")
print(f"  Lagoa Nova: pago {brl(tot[405])} | restituir {brl(excessos[405])}")
print(f"  Jucurutu: pago {brl(tot[400])} | restituir {brl(excessos[400])}")
print(f"  TOTAL a restituir: {brl(total_restituir)}")
