"""Informações de desconto em folha dos débitos de NEREU BATISTA LINHARES.

Dez processos de execução que (a) estão na CCD, (b) NÃO tratam de verba transitória
de servidor da SESAP e (c) já têm despacho da relatoria determinando o desconto em
folha (art. 118, I, da LCE nº 464/2012). Levantamento em
saidas/analise/nereu_ccd_despacho_desconto_folha.xlsx.

Valor = soma do valor atualizado (fn_Exe_RetornaValorAtualizado) dos débitos vigentes
do processo — cancelados ficam de fora. O órgão notificado é a SEAD, como nos despachos
de 06/05/2026 do relator nos processos 003023/2022 e 003659/2022.

Base: scripts/automacao/templates/desconto_folha.docx (modelo atualizado pelo usuário).
Rodar: .venv/Scripts/python.exe processos/gerar_informacoes_nereu_desconto_folha.py
"""
import shutil
from datetime import datetime
from pathlib import Path

from docxtpl import DocxTemplate
from num2words import num2words

from ccd.config import REPO_ROOT, cpf
from ccd.db import run_query_df
from ccd.docs import docx_to_pdf

TEMPLATE = str(REPO_ROOT / "scripts/automacao/templates/desconto_folha.docx")
DESTINO = Path(__file__).parent / "nereu_desconto_folha"

CPF = cpf("NEREU")
ORGAO = "SECRETARIA DE ADMINISTRAÇÃO DO ESTADO (SEAD)"
PROCESSOS = [  # 000130/2023 saiu: desconto já implementado
    "000101/2022", "000133/2022", "000142/2023", "002062/2024",
    "002564/2024", "003023/2022", "003064/2022", "003078/2022", "003659/2022",
]


def brl(valor: float) -> str:
    return "R$ " + f"{valor:,.2f}".replace(",", "@").replace(".", ",").replace("@", ".")


def por_extenso(valor: float) -> str:
    """'8985.10' -> 'oito mil, novecentos e oitenta e cinco reais e dez centavos'."""
    reais, centavos = f"{valor:.2f}".split(".")
    return (f"{num2words(int(reais), lang='pt_BR')} reais e "
            f"{num2words(int(centavos), lang='pt_BR')} centavos")


def carregar() -> list[dict]:
    """Débitos vigentes (não cancelados) de cada processo, com valor atualizado."""
    ph = {f"p{i}": p for i, p in enumerate(PROCESSOS)}
    inc = ", ".join(f":{k}" for k in ph)
    df = run_query_df(f"""
        SELECT CONCAT(RIGHT('000000'+CAST(pro.numero_processo AS varchar),6),'/',
                      pro.ano_processo) AS processo,
               RTRIM(pro.assunto) AS assunto, ed.IdDebito,
               RTRIM(s.DescricaoStatusDivida) AS situacao, gp.Nome AS nome,
               processo.dbo.fn_Exe_RetornaValorAtualizado(ed.IdDebito) AS valor
        FROM processo.dbo.Processos pro
        JOIN processo.dbo.Exe_Debito ed ON ed.IdProcessoExecucao = pro.IdProcesso
        JOIN processo.dbo.Exe_StatusDivida s
             ON s.CodigoStatusDivida = ed.CodigoStatusDivida
        JOIN processo.dbo.Exe_DebitoPessoa edp ON edp.IDDebito = ed.IdDebito
        JOIN processo.dbo.GenPessoa gp ON gp.IdPessoa = edp.IDPessoa
        WHERE gp.Documento = :cpf AND ed.DataCancelamento IS NULL
          AND CONCAT(RIGHT('000000'+CAST(pro.numero_processo AS varchar),6),'/',
                     pro.ano_processo) IN ({inc})""", cpf=CPF, **ph)

    itens = []
    for proc in PROCESSOS:
        g = df[df.processo == proc]
        assert not g.empty, f"{proc}: nenhum débito vigente do responsável"
        itens.append({
            "processo": proc,
            "assunto": g.assunto.iloc[0],
            "nome": g.nome.iloc[0],
            "valor": float(g.valor.sum()),
            "debitos": [(int(r.IdDebito), r.situacao) for r in g.itertuples()],
        })
    return itens


def contexto(item: dict) -> dict:
    return {
        "processo": item["processo"],
        "assunto": item["assunto"],
        "nome": f'{item["nome"]} (CPF: {CPF})',
        "orgao": ORGAO,
        "valor": f'{brl(item["valor"])} ({por_extenso(item["valor"])})',
    }


def gerar(item: dict) -> Path:
    num, ano = item["processo"].split("/")
    pasta = DESTINO / f"{num}_{ano}"
    pasta.mkdir(parents=True, exist_ok=True)
    out = pasta / f"informacao_{num}_{ano}.docx"
    doc = DocxTemplate(TEMPLATE)
    doc.render(contexto(item))
    # preserva a versão anterior antes de sobrescrever (skill edicao-minima)
    if out.exists():
        bak = out.with_name(f"{out.stem}_{datetime.now():%Y%m%d_%H%M%S}{out.suffix}")
        shutil.copy2(out, bak)
        print(f"  versão anterior preservada em: {bak.name}")
    doc.save(str(out))
    docx_to_pdf(str(out), str(pasta))
    return out


if __name__ == "__main__":
    import docx  # noqa: E402

    for item in carregar():
        ctx = contexto(item)
        deb = ", ".join(f"{i} ({s})" for i, s in item["debitos"])
        print(f'{item["processo"]}: {brl(item["valor"])} | débitos {deb}')
        out = gerar(item)

        # check (ponytail): nada de placeholder solto e os campos foram para o lugar certo
        d = docx.Document(str(out))
        texto = "\n".join([p.text for p in d.paragraphs]
                          + [c.text for t in d.tables for r in t.rows for c in r.cells])
        assert "{{" not in texto and "}}" not in texto, "placeholder não substituído"
        for campo, valor in ctx.items():
            assert valor in texto, f'{item["processo"]}: faltou {campo} ("{valor[:40]}")'
    print(f"\n{len(PROCESSOS)} informações em {DESTINO}")
